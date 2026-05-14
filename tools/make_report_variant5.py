# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Отчёт_мини_компилятор_вариант_5.docx"
GITHUB_BASE = "https://github.com/cergo9991/mini-compiler-variant-5/blob/codex/full-project-built"
GITHUB_SRC = f"{GITHUB_BASE}/src"


def set_run_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False, italic: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run_text = paragraph.add_run("1")
    set_run_font(run_text, size=12)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 13, bold=True)
    return p


def add_paragraph(doc: Document, text: str = "", first_indent: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.rstrip())
    set_run_font(run, name="Courier New", size=8)
    return p


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, center=True)
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_placeholder_box(doc: Document, title: str, command: str, expected: str):
    table = doc.add_table(rows=4, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fills = ["EAF2F8", "F8FBFF", "F8FBFF", "F8FBFF"]
    lines = [
        ("МЕСТО ДЛЯ СКРИНШОТА", True),
        (title, False),
        ("Что выполнить:\n" + command, False),
        ("Что должно быть видно:\n" + expected, False),
    ]
    for row, fill, (text, bold) in zip(table.rows, fills, lines):
        shade_cell(row.cells[0], fill)
        set_cell_text(row.cells[0], text, bold=bold, center=bold)
    doc.add_paragraph()


def add_toc(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = p.add_run("Оглавление будет заполнено после обновления полей в Word.")
    set_run_font(placeholder)

    run_end = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def add_title_page(doc: Document):
    lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "Псковский государственный университет",
        "Передовая инженерная школа гибридных технологий в станкостроении Союзного государства",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Курсовая работа на тему:")
    set_run_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Разработка компилятора для учебного языка программирования\nс генерацией MIPS-кода")
    set_run_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Вариант 5")
    set_run_font(run, size=14, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Выполнил\nстудент 3-го курса\nгр. 0483-04\nГерасимов В.А.")
    set_run_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Проверила\nТимошевская О. Ю.")
    set_run_font(run)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Псков 2026")
    set_run_font(run)

    doc.add_page_break()


def add_link_block(doc: Document, title: str, url: str):
    p = add_paragraph(doc, title, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    if p.runs:
        p.runs[0].bold = True
    p2 = add_paragraph(doc, url, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    return p2


def build_report():
    doc = Document()
    configure_document(doc)
    add_page_number(doc.sections[0])

    add_title_page(doc)

    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    add_toc(doc)
    add_paragraph(doc, "После открытия документа в Word обновите поля комбинацией Ctrl+A, затем F9.", first_indent=False)
    doc.add_page_break()

    add_heading(doc, "1. Задание", 1)
    add_paragraph(
        doc,
        "Компилятор в рамках данной работы понимается как программа-переводчик, которая принимает на вход текст "
        "на исходном языке и выдаёт эквивалентное представление на более низком уровне. Практически любой компилятор "
        "строится как последовательность фаз: лексический разбор, синтаксический разбор, семантическая проверка и "
        "генерация целевого кода."
    )
    add_paragraph(
        doc,
        "На вход разработанной системы подаётся файл с программой на учебном языке `.mc`, а на выходе формируются "
        "промежуточное представление LLVM IR и объектный файл. В данной работе все фазы реализованы раздельно, что "
        "упрощает проверку корректности и соответствует классической архитектуре компилятора."
    )
    add_paragraph(
        doc,
        "В рамках курсовой работы реализован учебный мини-компилятор, который выполняет полный цикл обработки "
        "программы: от чтения исходного текста до построения объектного файла для архитектуры MIPS и последующего "
        "запуска через внешнюю исполняющую среду на языке C."
    )
    add_paragraph(
        doc,
        "Требования задания выполнены в следующей конфигурации проекта: лексический анализ реализован на Flex, "
        "синтаксический анализ реализован на GNU Bison, семантическая проверка выделена в самостоятельную фазу, "
        "а генерация кода построена на инфраструктуре LLVM."
    )
    add_bullet(doc, "минимальный набор конструкций: арифметические и логические операции, условные операторы `if/else`, цикл `for`, функции и типы `int`, `float`.")
    add_bullet(doc, "стиль синтаксиса: C-style, с объявлениями функций в форме `int compiled_fn(int arg) { ... }` и блоками в фигурных скобках.")
    add_bullet(doc, "дополнительная конструкция варианта: цикл `do { ... } while (condition);`.")
    add_bullet(doc, "целевая архитектура: MIPS с целевой строкой `mipsel-unknown-linux-gnu`.")
    add_bullet(doc, "практический результат: генерация файлов `.ll` и `.o`, компоновка с `runtime/main.c` и запуск через `qemu-mipsel`.")
    add_paragraph(doc, "Основные этапы работы компилятора:")
    add_paragraph(doc, "1. Лексический анализ исходного файла выполняется на основе правил Flex.", first_indent=False)
    add_paragraph(doc, "2. GNU Bison строит абстрактное синтаксическое дерево, не смешивая синтаксический анализ с генерацией кода.", first_indent=False)
    add_paragraph(doc, "3. Семантический анализатор проверяет области видимости, типы, корректность `return`, вызовы функций и обязательную сигнатуру точки входа `compiled_fn`.", first_indent=False)
    add_paragraph(doc, "4. Модуль генерации кода преобразует AST в LLVM IR, а затем в объектный файл для MIPS.", first_indent=False)
    add_paragraph(doc, "5. Полученный объектный файл связывается с `runtime/main.c` и запускается в эмуляторе.", first_indent=False)
    add_paragraph(doc, "6. Корректность работы подтверждается набором тестов корректных программ, тестов ожидаемых ошибок и проверкой полного конвейера компиляции и запуска.", first_indent=False)

    add_heading(doc, "2. Лексическая спецификация", 1)
    add_paragraph(
        doc,
        "Лексический анализатор предназначен для преобразования потока символов в последовательность токенов, "
        "то есть элементарных единиц языка. На вход он получает исходный текст программы, а на выход передаёт "
        "синтаксическому анализатору токены, их значения и координаты в файле."
    )
    add_paragraph(
        doc,
        "В данном проекте эта фаза реализована в Flex. Правила описывают ключевые слова, литералы, идентификаторы, "
        "комментарии и операторы, а также обновляют сведения о строке и столбце для диагностики ошибок."
    )
    add_table(doc, ["Группа", "Описание"], [
        ["Ключевые слова", "`return`, `if`, `else`, `for`, `do`, `while`, `int`, `float`"],
        ["Идентификаторы", "Шаблон `[A-Za-z_][A-Za-z0-9_]*`"],
        ["Литералы", "`INTEGER`, `FLOAT_LITERAL`"],
        ["Операторы сравнения", "`==`, `!=`, `<=`, `>=`, `<`, `>`"],
        ["Логические операторы", "`&&`, `||`, `!`"],
        ["Арифметические операторы", "`+`, `-`, `*`, `/`, `%`"],
        ["Разделители", "`(`, `)`, `{`, `}`, `,`, `;`, `=`"],
    ])
    add_paragraph(
        doc,
        "Лексер также обрабатывает однострочные и многострочные комментарии и вычисляет координаты токенов через "
        "механизм `yylloc`, что позволяет выдавать точные сообщения вида `line X, column Y`."
    )

    add_heading(doc, "3. Диаграмма разработанных правил грамматики", 1)
    add_paragraph(
        doc,
        "Синтаксический анализ определяет, как из потока токенов строится корректная структура программы. "
        "Для текущего проекта грамматика должна отражать именно C-style объявления функций, выражения, блоки, "
        "конструкции `if/else`, `for`, `do/while`, вызовы функций и уровни приоритета операторов."
    )
    add_placeholder_box(
        doc,
        "Сюда вставляется диаграмма грамматики варианта 5.",
        "Вставить готовое изображение диаграммы вручную.",
        "На изображении должны быть показаны только реальные правила текущего проекта: `program`, `function_list`, "
        "`function`, `block`, `statement`, `expression`, `argument_list`, типы `int/float`, а также `if/else`, `for`, `do/while`, `return`."
    )

    add_heading(doc, "4. Листинг кода лексера и парсера", 1)
    add_paragraph(
        doc,
        "По аналогии с примером в основной части отчёта вместо полного листинга приведены ссылки на исходные файлы проекта "
        "в репозитории GitHub. Это позволяет сохранить компактность отчёта и при этом дать доступ к актуальному коду."
    )
    add_link_block(doc, "Файл `src/lexer.l`", f"{GITHUB_SRC}/lexer.l")
    add_link_block(doc, "Файл `src/parser.y`", f"{GITHUB_SRC}/parser.y")
    add_paragraph(
        doc,
        "Именно эти файлы реализуют front-end компилятора: `lexer.l` разбивает входной текст на токены, а `parser.y` "
        "описывает грамматику и строит AST."
    )

    add_heading(doc, "5. Описание семантических правил языка", 1)
    add_paragraph(
        doc,
        "Семантический анализ выполняет содержательную проверку программы после того, как её структура уже распознана. "
        "На вход этой фазе передаётся абстрактное синтаксическое дерево, а на выходе формируется либо подтверждение "
        "корректности программы, либо набор диагностических сообщений об ошибках."
    )
    add_paragraph(doc, "В текущем проекте реализованы следующие основные семантические правила:")
    add_paragraph(doc, "1. в программе обязательно должна существовать функция `compiled_fn`;", first_indent=False)
    add_paragraph(doc, "2. сигнатура обязательной функции должна быть строго `int compiled_fn(int arg)`;", first_indent=False)
    add_paragraph(doc, "3. в одной области видимости запрещено повторное объявление переменной;", first_indent=False)
    add_paragraph(doc, "4. использование необъявленного имени считается ошибкой;", first_indent=False)
    add_paragraph(doc, "5. арифметические операции разрешены только для согласованных числовых типов;", first_indent=False)
    add_paragraph(doc, "6. условие `if`, `for` и `do/while` обязано иметь тип `bool`;", first_indent=False)
    add_paragraph(doc, "7. проверяются типы выражения `return`, присваивания и аргументов при вызове функций.", first_indent=False)
    add_paragraph(
        doc,
        "Отдельно для варианта 5 важна проверка конструкции `do/while`: тело анализируется как отдельный блок, "
        "а выражение после `while` должно иметь логический тип `bool`."
    )
    add_paragraph(doc, "Пример семантической ошибки", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_placeholder_box(
        doc,
        "Семантическая ошибка: отсутствие обязательной точки входа",
        "./build/mini_cc tests/17_missing_compiled_fn.mc -o build/17_missing_compiled_fn.o",
        "На скриншоте должно быть видно сообщение `semantic error: required function 'compiled_fn' is missing`. "
        "Этот пример показывает, что семантический анализатор проверяет глобальную структуру всей программы, а не только локальные типы."
    )

    add_heading(doc, "6. Листинг кода семантического анализатора", 1)
    add_paragraph(
        doc,
        "Вместо полного листинга в основной части отчёта приводится ссылка на файл семантического анализатора в GitHub."
    )
    add_link_block(doc, "Файл `src/sema.cpp`", f"{GITHUB_SRC}/sema.cpp")
    add_paragraph(
        doc,
        "Именно в этом файле реализованы таблица функций, стек областей видимости, вывод типов выражений и текст "
        "диагностических сообщений об ошибках."
    )

    add_heading(doc, "7. Описание процесса перевода AST в IR и из IR в asm/obj", 1)
    add_paragraph(
        doc,
        "Генерация кода переводит внутреннее представление программы в форму, пригодную для дальнейшей машинной обработки. "
        "На вход эта фаза получает уже проверенное AST, а на выходе выдаёт промежуточное представление LLVM IR и объектный файл."
    )
    add_paragraph(
        doc,
        "В текущем проекте тип `int` учебного языка отображается в LLVM `i64`, тип `float` — в `double`, "
        "а служебный тип `bool` — в `i1`. Локальные переменные размещаются через `alloca`, чтение выполняется через `load`, "
        "а присваивание — через `store`."
    )
    add_paragraph(
        doc,
        "Для варианта 5 конструкция `do/while` переводится в три базовых блока: `do.body`, `do.cond` и `do.end`. "
        "Сначала управление безусловно переходит в тело цикла, затем выполняется проверка условия и либо происходит "
        "повторный переход в тело, либо выход из цикла."
    )
    add_code_block(
        doc,
        "entry:\n"
        "  br label %do.body\n\n"
        "do.body:\n"
        "  ; тело цикла\n"
        "  br label %do.cond\n\n"
        "do.cond:\n"
        "  ; вычисление условия\n"
        "  br i1 %condition, label %do.body, label %do.end\n\n"
        "do.end:\n"
        "  ; продолжение программы"
    )
    add_paragraph(
        doc,
        "После генерации IR модуль передаётся `TargetMachine`, настроенному на triple `mipsel-unknown-linux-gnu`. "
        "Результатом этой стадии становится объектный файл `.o`, который затем связывается с `runtime/main.c`."
    )

    add_heading(doc, "8. Листинг кода кодогенератора", 1)
    add_paragraph(
        doc,
        "Как и в случае с front-end и семантической фазой, вместо полного листинга приводится ссылка на файл "
        "кодогенератора в GitHub."
    )
    add_link_block(doc, "Файл `src/codegen.cpp`", f"{GITHUB_SRC}/codegen.cpp")
    add_paragraph(
        doc,
        "В этом файле реализованы объявление функций LLVM, генерация базовых блоков, short-circuit логика, "
        "обработка `do/while` и выпуск MIPS object file."
    )

    add_heading(doc, "9. Тестовые наборы программ с пояснениями", 1)
    add_paragraph(
        doc,
        "По структуре эта глава построена по образцу принятой работы: ниже приведены отдельные практические сценарии "
        "запуска тестов, а после каждого сценария оставлено место под соответствующий скриншот."
    )
    add_paragraph(
        doc,
        "Тесты охватывают как корректные программы, так и программы с ошибками разбора и семантики. "
        "Особое внимание уделено конструкциям варианта 5, то есть циклам `do/while`, а также целевой MIPS-генерации."
    )

    add_paragraph(doc, "Тест `05_do_while_sum`", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(
        doc,
        "Этот тест демонстрирует базовую работу конструкции `do/while`: программа суммирует числа от 0 до 9, "
        "поскольку `runtime/main.c` вызывает `compiled_fn(10)`."
    )
    add_paragraph(doc, "Команда 9.1 — запуск `05_do_while_sum` со стандартным аргументом", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_code_block(
        doc,
        "./build/mini_cc tests/05_do_while_sum.mc -o build/05_do_while_sum.o --emit-ir build/05_do_while_sum.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/05_do_while_sum.o -o build/05_do_while_sum.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/05_do_while_sum.mips"
    )
    add_placeholder_box(
        doc,
        "Запуск теста `05_do_while_sum`",
        "./build/mini_cc tests/05_do_while_sum.mc -o build/05_do_while_sum.o --emit-ir build/05_do_while_sum.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/05_do_while_sum.o -o build/05_do_while_sum.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/05_do_while_sum.mips",
        "На скриншоте должны быть видны успешная компиляция, линковка и итоговый вывод `45`. "
        "Этот результат подтверждает, что цикл `do/while` корректно отрабатывает несколько итераций и суммирует значения от 0 до 9."
    )

    add_paragraph(doc, "Команда 9.2 — повторный запуск `05_do_while_sum` с другим входным значением", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_code_block(
        doc,
        "cp runtime/main.c build/runtime_arg5.c\n"
        "sed -i 's/compiled_fn(10)/compiled_fn(5)/' build/runtime_arg5.c\n"
        "./build/mini_cc tests/05_do_while_sum.mc -o build/05_do_while_sum_arg5.o --emit-ir build/05_do_while_sum_arg5.ll\n"
        "mipsel-linux-gnu-gcc build/runtime_arg5.c build/05_do_while_sum_arg5.o -o build/05_do_while_sum_arg5.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/05_do_while_sum_arg5.mips"
    )
    add_placeholder_box(
        doc,
        "Повторный запуск `05_do_while_sum` с аргументом 5",
        "cp runtime/main.c build/runtime_arg5.c\n"
        "sed -i 's/compiled_fn(10)/compiled_fn(5)/' build/runtime_arg5.c\n"
        "./build/mini_cc tests/05_do_while_sum.mc -o build/05_do_while_sum_arg5.o --emit-ir build/05_do_while_sum_arg5.ll\n"
        "mipsel-linux-gnu-gcc build/runtime_arg5.c build/05_do_while_sum_arg5.o -o build/05_do_while_sum_arg5.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/05_do_while_sum_arg5.mips",
        "На скриншоте должен быть виден итоговый вывод `10`. Изменение результата подтверждает, что программа действительно использует "
        "входной аргумент, а не возвращает фиксированное значение."
    )

    add_paragraph(doc, "Тест `07_do_while_power`", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(
        doc,
        "Этот тест демонстрирует ещё один сценарий варианта 5, где `do/while` используется не для суммирования, "
        "а для поэтапного вычисления результата по нескольким итерациям."
    )
    add_paragraph(doc, "Команда 9.3 — запуск `07_do_while_power`", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_code_block(
        doc,
        "./build/mini_cc tests/07_do_while_power.mc -o build/07_do_while_power.o --emit-ir build/07_do_while_power.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/07_do_while_power.o -o build/07_do_while_power.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/07_do_while_power.mips"
    )
    add_placeholder_box(
        doc,
        "Запуск теста `07_do_while_power`",
        "./build/mini_cc tests/07_do_while_power.mc -o build/07_do_while_power.o --emit-ir build/07_do_while_power.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/07_do_while_power.o -o build/07_do_while_power.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/07_do_while_power.mips",
        "На скриншоте должен быть виден итоговый вывод `16`. Этот пример показывает, что цикл `do/while` корректно поддерживает "
        "многократное изменение промежуточного результата."
    )

    add_paragraph(doc, "Тест `12_float_arithmetic`", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(
        doc,
        "Этот тест полезен тем, что одновременно проверяет поддержку типа `float`, арифметику над `float`, сравнение и совместную "
        "работу `float` с конструкцией `do/while`."
    )
    add_paragraph(doc, "Команда 9.4 — запуск `12_float_arithmetic`", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_code_block(
        doc,
        "./build/mini_cc tests/12_float_arithmetic.mc -o build/12_float_arithmetic.o --emit-ir build/12_float_arithmetic.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/12_float_arithmetic.o -o build/12_float_arithmetic.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/12_float_arithmetic.mips"
    )
    add_placeholder_box(
        doc,
        "Запуск теста `12_float_arithmetic`",
        "./build/mini_cc tests/12_float_arithmetic.mc -o build/12_float_arithmetic.o --emit-ir build/12_float_arithmetic.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/12_float_arithmetic.o -o build/12_float_arithmetic.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/12_float_arithmetic.mips",
        "На скриншоте должен быть виден итоговый вывод `12`. Этот результат подтверждает, что поддержка `float` "
        "работает не только на уровне синтаксиса, но и на уровне семантики и генерации MIPS-кода."
    )

    add_heading(doc, "10. Описание процесса использования компилятора", 1)
    add_paragraph(
        doc,
        "Практическое использование компилятора сводится к трём основным действиям: сборке проекта, компиляции исходной "
        "программы и запуску полученного MIPS-бинарника через эмулятор."
    )
    add_paragraph(
        doc,
        "В текущем проекте рабочий сценарий построен вокруг CMake, исполняемого файла `mini_cc`, компоновки через "
        "`mipsel-linux-gnu-gcc` и запуска в `qemu-mipsel`."
    )
    add_paragraph(doc, "Сборка проекта через CMake:", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Команда 10.1", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_code_block(
        doc,
        "cmake -S . -B build -DLLVM_DIR=/usr/lib/llvm-14/cmake\n"
        "cmake --build build -j"
    )
    add_paragraph(doc, "Компиляция тестовой программы с генерацией `.ll` и `.o`:", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Команда 10.2", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_code_block(
        doc,
        "./build/mini_cc tests/05_do_while_sum.mc -o build/05_do_while_sum.o --emit-ir build/05_do_while_sum.ll"
    )
    add_paragraph(doc, "Линковка с `runtime/main.c` и запуск в `qemu-mipsel`:", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Команда 10.3", first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_code_block(
        doc,
        "mipsel-linux-gnu-gcc runtime/main.c build/05_do_while_sum.o -o build/05_do_while_sum.mips\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/05_do_while_sum.mips"
    )

    add_heading(doc, "11. Результаты тестирования", 1)
    add_paragraph(
        doc,
        "Результаты тестирования показывают, насколько реализованный компилятор соответствует заявленным требованиям не только "
        "теоретически, но и практически. Здесь важны три группы результатов: корректность генерации LLVM IR и MIPS object file, "
        "успешное выполнение положительных сценариев и корректная остановка на отрицательных тестах."
    )
    add_paragraph(doc, "На момент подготовки отчёта зафиксированы следующие результаты:", first_indent=False)
    add_paragraph(doc, "• `05_do_while_sum`: вывод `45`;", first_indent=False)
    add_paragraph(doc, "• `07_do_while_power`: вывод `16`;", first_indent=False)
    add_paragraph(doc, "• `12_float_arithmetic`: вывод `12`;", first_indent=False)
    add_paragraph(doc, "• `14_parse_missing_semicolon`: корректный `parse error`;", first_indent=False)
    add_paragraph(doc, "• `17_missing_compiled_fn`: корректный `semantic error`;", first_indent=False)
    add_paragraph(doc, "Полный контрольный прогон выбранных тестов:", first_indent=False)
    add_code_block(
        doc,
        "./build/mini_cc tests/05_do_while_sum.mc -o build/05_do_while_sum.o --emit-ir build/05_do_while_sum.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/05_do_while_sum.o -o build/05_do_while_sum.mips\n"
        "printf '05_do_while_sum: '\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/05_do_while_sum.mips\n\n"
        "./build/mini_cc tests/07_do_while_power.mc -o build/07_do_while_power.o --emit-ir build/07_do_while_power.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/07_do_while_power.o -o build/07_do_while_power.mips\n"
        "printf '07_do_while_power: '\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/07_do_while_power.mips\n\n"
        "./build/mini_cc tests/12_float_arithmetic.mc -o build/12_float_arithmetic.o --emit-ir build/12_float_arithmetic.ll\n"
        "mipsel-linux-gnu-gcc runtime/main.c build/12_float_arithmetic.o -o build/12_float_arithmetic.mips\n"
        "printf '12_float_arithmetic: '\n"
        "qemu-mipsel -L /usr/mipsel-linux-gnu build/12_float_arithmetic.mips\n\n"
        "./build/mini_cc tests/14_parse_missing_semicolon.mc -o build/14_parse_missing_semicolon.o\n"
        "./build/mini_cc tests/17_missing_compiled_fn.mc -o build/17_missing_compiled_fn.o"
    )
    add_placeholder_box(
        doc,
        "Контрольный прогон тестов",
        "Запустить команды из блока выше последовательно в WSL Ubuntu.",
        "На скриншоте должны быть видны результаты `45`, `16`, `12`, а также сообщения `parse error` и `semantic error`. "
        "Такой снимок подтверждает, что положительные тесты проходят полный конвейер, а отрицательные корректно останавливаются на нужной фазе."
    )

    add_heading(doc, "12. Заключение", 1)
    add_paragraph(
        doc,
        "В ходе выполнения курсовой работы был разработан учебный мини-компилятор для языка программирования "
        "с C-style синтаксисом и генерацией кода под архитектуру MIPS. Реализованный проект охватывает основные этапы "
        "компиляции: лексический, синтаксический и семантический анализ, построение AST, генерацию LLVM IR и выпуск MIPS object file."
    )
    add_paragraph(
        doc,
        "Особое требование варианта 5 — цикл `do/while` — реализовано полностью: поддерживаются его лексический и синтаксический разбор, "
        "AST-узел, семантическая проверка условия и генерация соответствующих базовых блоков `do.body`, `do.cond` и `do.end` в LLVM IR."
    )
    add_paragraph(
        doc,
        "Результаты тестирования показывают, что компилятор корректно работает как на допустимых программах, так и на ошибочных входных данных, "
        "а связка `mini_cc` -> MIPS object file -> `runtime/main.c` -> `qemu-mipsel` образует завершённый практический конвейер."
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_report()
    print(OUT)
