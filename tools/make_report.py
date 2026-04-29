from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Отчёт_мини_компилятор_вариант_11.docx"
GUIDE = ROOT / "Инструкция_скриншоты.md"
DIAGRAM = ROOT / "diagram_grammar_variant_11.png"
POSITIVE_TESTS_COMMAND = """cd "/mnt/c/Users/minae/Documents/New project/build"
for t in ../tests/01_return_literal.mc ../tests/02_arithmetic.mc ../tests/03_if_else.mc ../tests/04_for.mc ../tests/05_ternary_true.mc ../tests/06_ternary_false.mc ../tests/07_nested_ternary.mc ../tests/10_short_circuit.mc ../tests/11_function_call.mc ../tests/test.mc; do
  base=$(basename "$t" .mc)
  ./mini_cc "$t" -o "$base.o" --emit-ir "$base.ll"
  clang ../runtime/main.c "$base.o" -o "$base.app"
  printf "%s: " "$base"
  ./"$base.app"
done"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc: Document, text: str = "", bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_code(doc: Document, text: str, max_lines: int | None = None) -> None:
    lines = text.rstrip().splitlines()
    if max_lines is not None and len(lines) > max_lines:
        lines = lines[:max_lines] + ["// ... фрагмент сокращён, полный файл находится в проекте"]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)


def add_placeholder(doc: Document, title: str, command: str, where: str) -> None:
    table = doc.add_table(rows=3, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, "EAF2F8")
    set_cell_text(table.cell(0, 0), f"МЕСТО ДЛЯ СКРИНШОТА: {title}", True)
    set_cell_text(table.cell(1, 0), f"Что должно быть видно: {where}")
    set_cell_text(table.cell(2, 0), f"Команда/действие: {command}")
    doc.add_paragraph()


def add_picture(doc: Document, path: Path, caption: str, width_cm: float = 16.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.name = "Times New Roman"
    cap.runs[0].font.size = Pt(11)
    cap.runs[0].italic = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        set_cell_shading(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def read(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Times New Roman"
        styles[name].font.color.rgb = RGBColor(31, 78, 121)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill=(0, 0, 0)) -> None:
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=3, align="center")
    x = box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2
    y = box[1] + (box[3] - box[1] - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=3, align="center")


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill=(25, 25, 25)) -> None:
    draw.line([start, end], fill=fill, width=2)
    x, y = end
    draw.polygon([(x, y), (x - 10, y - 5), (x - 10, y + 5)], fill=fill)


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, kind: str, font) -> None:
    if kind == "terminal":
        fill = (255, 255, 255)
        outline = (185, 55, 55)
        text_fill = (185, 0, 0)
    elif kind == "nonterminal":
        fill = (224, 236, 255)
        outline = (66, 103, 178)
        text_fill = (25, 50, 120)
    else:
        fill = (255, 255, 255)
        outline = (40, 40, 40)
        text_fill = (20, 20, 20)
    draw.rounded_rectangle(box, radius=8, fill=fill, outline=outline, width=2)
    draw_centered(draw, box, text, font, text_fill)


def draw_rule(draw: ImageDraw.ImageDraw, x: int, y: int, title: str, items: list[tuple[str, str]], font, small) -> None:
    panel = (x, y, x + 570, y + 105)
    draw.rounded_rectangle(panel, radius=10, fill=(255, 255, 255), outline=(60, 60, 60), width=2)
    draw.text((x + 12, y + 8), f"{title} =", font=small, fill=(0, 0, 0))
    line_y = y + 60
    draw_arrow(draw, (x + 22, line_y), (x + 58, line_y))
    cur = x + 65
    for label, kind in items:
        w = max(58, int(draw.textlength(label, font=font)) + 24)
        box = (cur, line_y - 17, cur + w, line_y + 17)
        draw_box(draw, box, label, kind, font)
        cur += w
        draw_arrow(draw, (cur, line_y), (cur + 35, line_y))
        cur += 42


def create_grammar_diagram() -> None:
    img = Image.new("RGB", (1800, 1320), "white")
    draw = ImageDraw.Draw(img)
    title = load_font(44, True)
    subtitle = load_font(22)
    font = load_font(18)
    small = load_font(16)
    tiny = load_font(14)
    heading = load_font(20, True)

    draw.text((620, 30), "MINI_CC parser grammar", font=title, fill=(0, 0, 0))
    draw.text((710, 82), "Railroad-style diagram generated for variant 11", font=subtitle, fill=(80, 80, 80))

    left_rules = [
        ("program", [("function_list", "nonterminal")]),
        ("function_list", [("function", "nonterminal")]),
        ("function", [("fn", "terminal"), ("IDENT", "terminal"), ("signature", "nonterminal"), ("block", "nonterminal")]),
        ("params", [("empty", "plain"), ("param_list", "nonterminal")]),
        ("param", [("int", "terminal"), ("IDENT", "terminal")]),
        ("block", [("{", "terminal"), ("statement_list", "nonterminal"), ("}", "terminal")]),
        ("var_decl", [("int", "terminal"), ("IDENT", "terminal"), ("=", "terminal"), ("expression", "nonterminal")]),
    ]
    for i, (name, items) in enumerate(left_rules):
        draw_rule(draw, 40, 130 + i * 118, name, items, font, small)

    mid_rules = [
        ("statement", [("var_decl", "nonterminal"), (";", "terminal")]),
        ("statement", [("expression", "nonterminal"), (";", "terminal")]),
        ("statement", [("return", "terminal"), ("expression", "nonterminal"), (";", "terminal")]),
        ("if_stmt", [("if", "terminal"), ("condition", "nonterminal"), ("then", "nonterminal"), ("else?", "nonterminal")]),
        ("for_stmt", [("for", "terminal"), ("for_header", "nonterminal"), ("block", "nonterminal")]),
        ("assign", [("IDENT", "terminal"), ("=", "terminal"), ("expression", "nonterminal")]),
        ("call", [("IDENT", "terminal"), ("(", "terminal"), ("args", "nonterminal"), (")", "terminal")]),
    ]
    for i, (name, items) in enumerate(mid_rules):
        draw_rule(draw, 625, 130 + i * 118, name, items, font, small)

    expr_panel = (1210, 130, 1760, 950)
    draw.rounded_rectangle(expr_panel, radius=12, fill=(255, 255, 255), outline=(60, 60, 60), width=2)
    draw.text((1230, 150), "expression =", font=heading, fill=(0, 0, 0))
    expression_lines = [
        ("Primary", ["INTEGER", "IDENT", "call", "( expression )"]),
        ("Unary", ["- expression", "! expression"]),
        ("Multiplicative", ["expression * expression", "expression / expression", "expression % expression"]),
        ("Additive", ["expression + expression", "expression - expression"]),
        ("Comparison", ["==", "!=", "<", "<=", ">", ">="]),
        ("Logical", ["expression && expression", "expression || expression"]),
        ("Ternary", ["condition ? then_expr : else_expr"]),
    ]
    positions = [
        (1230, 195),
        (1230, 400),
        (1230, 550),
        (1230, 775),
        (1500, 195),
        (1500, 470),
        (1500, 620),
    ]
    for (group, lines), (gx, gy) in zip(expression_lines, positions):
        draw.text((gx, gy), group, font=small, fill=(0, 0, 0))
        y = gy + 30
        for line in lines:
            draw_box(draw, (gx + 15, y, gx + 235, y + 34), line, "nonterminal" if "expression" in line or "expr" in line else "terminal", tiny)
            y += 42

    prec = (1210, 980, 1760, 1300)
    draw.rounded_rectangle(prec, radius=12, fill=(255, 255, 255), outline=(60, 60, 60), width=2)
    draw.text((1360, 995), "Operator precedence", font=heading, fill=(0, 0, 0))
    rows = [
        ("1", "?:", "right"),
        ("2", "||", "left"),
        ("3", "&&", "left"),
        ("4", "== !=", "left"),
        ("5", "< <= > >=", "left"),
        ("6", "+ -", "left"),
        ("7", "* / %", "left"),
        ("8", "! unary -", "right"),
    ]
    x0, y0 = 1235, 1040
    col_w = [50, 220, 180]
    for i, row in enumerate(rows):
        yy = y0 + i * 24
        draw.text((x0, yy), row[0], font=tiny, fill=(0, 0, 0))
        draw.text((x0 + col_w[0], yy), row[1], font=tiny, fill=(190, 0, 0))
        draw.text((x0 + col_w[0] + col_w[1], yy), row[2], font=tiny, fill=(0, 0, 0))

    token_panel = (40, 980, 1170, 1300)
    draw.rounded_rectangle(token_panel, radius=12, fill=(255, 255, 255), outline=(60, 60, 60), width=2)
    draw.text((500, 995), "Parser tokens", font=heading, fill=(0, 0, 0))
    tokens = [
        ("Keywords", ["fn", "int", "if", "else", "for", "return"]),
        ("Identifiers and literals", ["IDENT", "INTEGER"]),
        ("Operators", ["=", "+", "-", "*", "/", "%", "==", "!=", "<", "<=", ">", ">=", "&&", "||", "!", "?", ":"]),
        ("Separators", ["(", ")", "{", "}", ",", ";"]),
    ]
    group_x = [70, 410, 750, 1040]
    group_limit = [270, 270, 235, 120]
    for (group, vals), tx, limit in zip(tokens, group_x, group_limit):
        draw.text((tx, 1035), group, font=small, fill=(0, 0, 0))
        xx = tx
        yy = 1070
        for val in vals:
            w = 54 if len(val) <= 2 else 64
            draw_box(draw, (xx, yy, xx + w, yy + 34), val, "terminal", tiny)
            xx += w + 10
            if xx > tx + limit:
                xx = tx
                yy += 44

    legend = (1470, 70, 1760, 115)
    draw_box(draw, (1485, 78, 1615, 108), "nonterminal", "nonterminal", tiny)
    draw_box(draw, (1625, 78, 1745, 108), "terminal", "terminal", tiny)

    img.save(DIAGRAM)


def make_report() -> None:
    create_grammar_diagram()
    doc = Document()
    configure_doc(doc)

    # Title page.
    for text in [
        "Министерство науки и высшего образования Российской Федерации",
        "Псковский государственный университет",
        "Передовая инженерная школа",
        "гибридных технологий в станкостроении",
        "Союзного государства",
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph("Курсовая работа на тему:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True

    p = doc.add_paragraph("Разработка мини-компилятора учебного языка программирования с генерацией объектного файла x86_64")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True

    p = doc.add_paragraph("Вариант 11")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph("Выполнил\nстудент 3-го курса\nгр. 0483-04\nМинаев А.Н.")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(12)
    p = doc.add_paragraph("Проверила\nТимошевская О. Ю.")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(12)
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph("Псков 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(12)

    doc.add_page_break()
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    toc_items = [
        "1. Задание",
        "2. Лексическая спецификация",
        "3. Диаграмма и описание правил грамматики",
        "4. Листинг кода лексического и синтаксического анализаторов",
        "5. Описание семантических правил языка",
        "6. Листинг кода семантического анализатора",
        "7. Перевод AST в LLVM IR и объектный файл",
        "8. Листинг кода генератора кода",
        "9. Тестовые наборы программ",
        "10. Описание процесса использования компилятора",
        "11. Результаты тестирования",
        "12. Заключение",
        "13. Приложения",
    ]
    for item in toc_items:
        add_para(doc, item)

    doc.add_page_break()
    add_heading(doc, "1. Задание", 1)
    add_para(doc, "В рамках курсовой работы разработан мини-компилятор учебного языка программирования. Компилятор принимает на вход файл с программой на языке .mc, выполняет лексический и синтаксический анализ, строит абстрактное синтаксическое дерево, проводит семантическую проверку и генерирует объектный файл для Linux.")
    add_para(doc, "Вариант 11 требует реализации тернарного оператора cond ? a : b. Целевая архитектура — x86_64 Linux, целевая строка LLVM — x86_64-pc-linux-gnu. Проект рассчитан на среду Ubuntu WSL, C++17 и LLVM/Clang версии 14.0.0.")
    add_table(doc, ["Параметр", "Значение"], [
        ["Язык реализации", "C++17"],
        ["Лексический анализ", "Flex"],
        ["Синтаксический анализ", "GNU Bison"],
        ["Промежуточное представление", "LLVM IR"],
        ["LLVM API", "LLVM 14, legacy::PassManager"],
        ["Целевая архитектура", "x86_64-pc-linux-gnu"],
        ["Генерируемая функция", "extern \"C\" int64_t compiled_fn(int64_t arg)"],
    ])

    add_heading(doc, "2. Лексическая спецификация", 1)
    add_para(doc, "Лексический анализ — это первая фаза компиляции, на которой непрерывный текст программы разбивается на отдельные смысловые единицы: ключевые слова, идентификаторы, литералы, операторы и разделители. На вход этой фазы поступает исходный файл .mc, а на выходе получается поток токенов, пригодный для синтаксического анализа.")
    add_para(doc, "На практике лексический анализ обычно реализуется конечным автоматом или генератором лексеров. В данной работе используется Flex — генератор лексических анализаторов, который по набору регулярных выражений создаёт C/C++-код функции сканирования.")
    add_para(doc, "Лексический анализатор преобразует исходный текст программы в поток токенов. Он распознаёт ключевые слова, идентификаторы, целочисленные литералы, операторы, скобки и комментарии. Для диагностики ошибок лексер также передаёт координаты токенов: строку и столбец.")
    add_para(doc, "В данном проекте лексер находится в файле src/lexer.l. Он возвращает токены в parser, сохраняет значения целочисленных литералов и идентификаторов, пропускает пробельные символы и комментарии, а также поддерживает координаты для последующих сообщений об ошибках.")
    add_table(doc, ["Группа", "Элементы"], [
        ["Ключевые слова", "fn, int, if, else, for, return"],
        ["Операторы", "+, -, *, /, %, ==, !=, <, <=, >, >=, &&, ||, !, ?, :, ="],
        ["Разделители", "(, ), {, }, ,, ;"],
        ["Комментарии", "однострочные // и многострочные /* ... */"],
        ["Идентификаторы", "[A-Za-z_][A-Za-z0-9_]*"],
        ["Литералы", "целые числа, компилируемые в i64"],
    ])

    add_heading(doc, "3. Диаграмма и описание правил грамматики", 1)
    add_para(doc, "Синтаксический анализ — это фаза, которая проверяет, соответствует ли последовательность токенов правилам грамматики языка. На вход поступает поток токенов от лексера, а на выходе формируется абстрактное синтаксическое дерево, то есть древовидное представление программы без лишних деталей записи.")
    add_para(doc, "На практике синтаксический анализ реализуют вручную или с помощью генераторов парсеров. В работе используется GNU Bison — генератор синтаксических анализаторов. Он принимает описание грамматики и создаёт parser, который вызывает действия при распознавании правил.")
    add_para(doc, "Синтаксический анализатор построен на GNU Bison. Его задача — проверить структуру программы и построить AST. На этом этапе LLVM IR не создаётся, что сохраняет разделение фаз компилятора.")
    add_para(doc, "AST — абстрактное синтаксическое дерево, то есть внутренняя структура программы после parsing. LLVM IR — промежуточное представление LLVM, но оно появляется только на этапе генерации кода, а не в grammar-файле parser.y.")
    add_picture(doc, DIAGRAM, "Рисунок 3.1 — диаграмма грамматики учебного языка варианта 11", width_cm=16.5)
    add_code(doc, """program
  -> function_list

function
  -> fn IDENT '(' param_list ')' -> type block

statement
  -> var_decl ';'
  -> expression ';'
  -> return expression ';'
  -> if '(' expression ')' block else block
  -> for '(' init ';' condition ';' step ')' block

expression
  -> assignment
  -> ternary
  -> logical_or '?' expression ':' ternary
  -> arithmetic / comparison / logical operators""")

    add_heading(doc, "4. Листинг кода лексического и синтаксического анализаторов", 1)
    add_para(doc, "Листинг front-end части показывает, как описание языка разделено между lexer и parser. Front-end, или передняя часть компилятора, отвечает за чтение исходной программы и построение AST. Входом является текст программы, а выходом — дерево узлов выражений, операторов и функций.")
    add_para(doc, "В проекте lexer.l содержит регулярные выражения для токенов, а parser.y содержит правила построения AST. Важно, что в parser.y отсутствует генерация LLVM-кода: синтаксический анализ только создаёт объекты AST, которые затем передаются в semantic analysis и code generation.")
    add_para(doc, "Ниже приведены ключевые фрагменты Flex и Bison. Полные файлы вынесены в приложения.")
    add_heading(doc, "Листинг 4.1 — фрагмент src/lexer.l", 2)
    add_code(doc, read("src/lexer.l"), max_lines=45)
    add_heading(doc, "Листинг 4.2 — фрагмент src/parser.y", 2)
    add_code(doc, read("src/parser.y"), max_lines=90)

    add_heading(doc, "5. Описание семантических правил языка", 1)
    add_para(doc, "Семантический анализ — это фаза компилятора, которая проверяет смысловую корректность уже разобранной программы. Синтаксически правильная программа всё ещё может быть ошибочной: например, она может использовать необъявленную переменную или возвращать значение неправильного типа.")
    add_para(doc, "На вход семантического анализатора поступает AST. На выходе получается либо подтверждение корректности программы, либо список диагностических сообщений. Кроме того, каждому выражению присваивается вычисленный тип, который затем используется генератором кода.")
    add_para(doc, "На практике semantic analysis строится вокруг таблиц символов и областей видимости. В данном проекте используется стек таблиц символов для локальных переменных и отдельная таблица функций. Такой подход позволяет проверять вложенные блоки, параметры функций и повторные объявления.")
    add_para(doc, "Семантический анализатор работает поверх готового AST. Он поддерживает глобальную таблицу функций и стек локальных областей видимости. Каждое выражение получает вычисленный тип, который затем используется генератором кода.")
    add_table(doc, ["Правило", "Описание"], [
        ["Объявления", "Повторное объявление имени в одной области видимости запрещено."],
        ["Использование имён", "Использование необъявленной переменной является ошибкой."],
        ["Тип int", "Единственный пользовательский тип компилируется в LLVM i64."],
        ["Bool выражения", "Результаты сравнений и логических операций имеют тип bool и компилируются в i1."],
        ["Условия", "if, for, &&, ||, ! и условие тернарного оператора требуют bool."],
        ["Return", "Функция compiled_fn обязана возвращать int."],
        ["Ternary", "Условие имеет тип bool, ветви имеют одинаковый тип."],
    ])

    add_heading(doc, "6. Листинг кода семантического анализатора", 1)
    add_para(doc, "Код семантического анализатора реализует обход AST. Для каждого узла выбирается проверка, соответствующая его виду: объявление переменной добавляется в текущую область видимости, обращение к переменной ищется в стеке областей, выражения проверяются по правилам типизации.")
    add_para(doc, "Входом функций SemanticAnalyzer является дерево Program. Выходом является булево значение успешности анализа и список ошибок. Если ошибки найдены, генерация LLVM IR не запускается, потому что дальнейшие фазы работают только с корректной программой.")
    add_para(doc, "Ключевая часть семантического анализатора показывает регистрацию функций, проверку областей видимости, типов и return. Диагностика выводится с координатами строки и столбца.")
    add_placeholder(doc, "Семантические ошибки", "./mini_cc ../tests/08_type_error.mc -o bad.o --emit-ir bad.ll\n./mini_cc ../tests/09_undeclared_variable.mc -o bad.o --emit-ir bad.ll", "semantic error с номером строки и столбца.")
    add_code(doc, read("src/sema.cpp"), max_lines=120)

    add_heading(doc, "7. Перевод AST в LLVM IR и объектный файл", 1)
    add_para(doc, "Генерация кода — это фаза, преобразующая проверенное AST в более низкоуровневое представление. В данном проекте сначала строится LLVM IR, то есть промежуточное представление инфраструктуры LLVM. Затем LLVM переводит это представление в объектный файл для архитектуры x86_64 Linux.")
    add_para(doc, "На вход генератора кода поступает AST после успешной семантической проверки. На выходе создаются файл LLVM IR с расширением .ll и объектный файл .o. Объектный файл затем связывается с runtime/main.c компилятором clang.")
    add_para(doc, "На практике LLVM IR строится через LLVMContext, Module и IRBuilder. LLVMContext хранит глобальное состояние LLVM, Module представляет единицу компиляции, а IRBuilder помогает создавать инструкции. В проекте эти объекты используются в src/codegen.cpp.")
    add_para(doc, "Генератор кода преобразует AST в LLVM IR. Для работы с переменными используется alloca в entry block функции, а обращения к переменным выполняются через load/store. Ручная SSA-форма не строится.")
    add_para(doc, "Для тернарного оператора создаются три basic block: then, else и merge. После вычисления двух ветвей в merge-блоке создаётся PHI node, выбирающий значение в зависимости от фактически выполненной ветви.")
    add_para(doc, "Объектный файл создаётся средствами LLVM 14: InitializeAllTargets, TargetRegistry::lookupTarget, createTargetMachine, legacy::PassManager и старая форма addPassesToEmitFile(..., CGFT_ObjectFile). Новый PassBuilder не используется.")
    add_placeholder(doc, "Проверка PHI для ternary", "grep -n phi generated.ll", "Строки LLVM IR с инструкцией phi, подтверждающие генерацию тернарного оператора через PHI node.")
    add_code(doc, """LLVMContext context;
Module module("mini_module", context);
IRBuilder<> builder(context);

TargetRegistry::lookupTarget("x86_64-pc-linux-gnu", error);
target->createTargetMachine(...);
legacy::PassManager pass;
targetMachine->addPassesToEmitFile(pass, dest, nullptr, CGFT_ObjectFile);""")

    add_heading(doc, "8. Листинг кода генератора кода", 1)
    add_para(doc, "Листинг генератора кода показывает практическую реализацию обхода AST и построения LLVM IR. Для каждого типа узла есть соответствующая ветка: литералы превращаются в константы, переменные — в load/store, арифметика — в инструкции сложения, вычитания, умножения, деления и остатка.")
    add_para(doc, "Отдельно реализованы logical operations, то есть логические операции, с коротким замыканием. Короткое замыкание означает, что правая часть выражения && или || вычисляется только тогда, когда это действительно необходимо. В LLVM IR это реализуется через basic blocks и PHI node.")
    add_para(doc, "В листинге показаны фрагменты генерации LLVM IR и объектного файла. Полный файл находится в приложении.")
    add_code(doc, read("src/codegen.cpp"), max_lines=160)

    add_heading(doc, "9. Тестовые наборы программ", 1)
    add_para(doc, "Тестирование компилятора проверяет не только отдельные функции программы, но и весь конвейер: исходный файл .mc, разбор, семантический анализ, генерацию LLVM IR, создание объектного файла, линковку с runtime и запуск итогового приложения.")
    add_para(doc, "На вход каждого теста подаётся отдельная программа учебного языка. На выходе ожидается либо корректное число после запуска скомпилированной программы, либо диагностическое сообщение semantic error для ошибочных программ.")
    tests = [
        ["01_return_literal.mc", "Возврат литерала"],
        ["02_arithmetic.mc", "Арифметика + - * / %"],
        ["03_if_else.mc", "Условная конструкция if/else"],
        ["04_for.mc", "Цикл for"],
        ["05_ternary_true.mc", "Тернарный оператор, true-ветвь"],
        ["06_ternary_false.mc", "Тернарный оператор, false-ветвь"],
        ["07_nested_ternary.mc", "Вложенный тернарный оператор"],
        ["08_type_error.mc", "Ошибка типов"],
        ["09_undeclared_variable.mc", "Необъявленная переменная"],
        ["10_short_circuit.mc", "Short-circuit для &&"],
        ["11_function_call.mc", "Вызов пользовательской функции"],
    ]
    add_table(doc, ["Файл", "Назначение"], tests)
    add_heading(doc, "Пример входной программы", 2)
    add_code(doc, read("tests/test.mc"))

    add_heading(doc, "10. Описание процесса использования компилятора", 1)
    add_para(doc, "Использование компилятора состоит из трёх этапов: сборка самого mini_cc, компиляция учебной программы в объектный файл и связывание объектного файла с runtime. Runtime — это небольшой C-файл, который вызывает сгенерированную функцию compiled_fn.")
    add_para(doc, "CMake — система сборки, которая генерирует файлы сборки для проекта. В данной работе CMake находит Flex, Bison и LLVM 14, затем собирает исполняемый файл mini_cc.")
    add_para(doc, "Сборка проекта выполняется стандартным CMake-процессом в Ubuntu WSL.")
    add_placeholder(doc, "Структура проекта", "cd \"/mnt/c/Users/minae/Documents/New project\"\nls\nfind src runtime tests -maxdepth 1 -type f | sort", "Папки src, runtime, tests, CMakeLists.txt, README.md и файлы исходного кода.")
    add_code(doc, """cd "/mnt/c/Users/minae/Documents/New project"
mkdir -p build
cd build
cmake ..
cmake --build .""")
    add_placeholder(doc, "Успешная сборка CMake", "mkdir -p build && cd build && cmake .. && cmake --build .", "Сообщение Found LLVM 14.0.0 и строка [100%] Built target mini_cc.")
    add_para(doc, "После сборки компилятор запускается командой:")
    add_code(doc, """./mini_cc ../tests/test.mc -o generated.o --emit-ir generated.ll
clang ../runtime/main.c generated.o -o app
./app""")
    add_placeholder(doc, "Запуск основного примера", "./mini_cc ../tests/test.mc -o generated.o --emit-ir generated.ll\nclang ../runtime/main.c generated.o -o app\n./app", "Вывод программы: 14.")

    add_heading(doc, "11. Результаты тестирования", 1)
    add_para(doc, "Результаты тестирования показывают, что реализованные фазы компилятора работают согласованно. Положительные тесты проходят полный путь до запуска исполняемого файла, а отрицательные тесты останавливаются на семантической фазе.")
    add_para(doc, "Особенно важны тесты тернарного оператора и short-circuit логики. Они подтверждают требования варианта: тернарный оператор выбирает одну из двух ветвей через PHI node, а логические операции не вычисляют лишние выражения.")
    add_para(doc, "Ниже приведён ожидаемый результат выполнения положительных тестов после компиляции, линковки с runtime/main.c и запуска полученной программы.")
    add_placeholder(doc, "Прогон положительных тестов", POSITIVE_TESTS_COMMAND, "Вывод всех положительных тестов: 42, 35, 11, 45, 100, 200, 1, 123, 15, 14.")
    add_code(doc, """01_return_literal: 42
02_arithmetic: 35
03_if_else: 11
04_for: 45
05_ternary_true: 100
06_ternary_false: 200
07_nested_ternary: 1
10_short_circuit: 123
11_function_call: 15
test: 14""")
    add_para(doc, "Ошибочные тесты должны завершаться semantic error:")
    add_code(doc, """status_08=1
semantic error: line 2, column 5: initializer for variable 'x' has type bool, expected int
status_09=1
semantic error: line 2, column 12: use of undeclared variable 'missing'""")
    add_placeholder(doc, "Демонстрация short-circuit", "./mini_cc ../tests/10_short_circuit.mc -o sc.o --emit-ir sc.ll\nclang ../runtime/main.c sc.o -o sc_app\n./sc_app", "Вывод 123; деление на ноль во второй части && не выполняется.")

    add_heading(doc, "12. Заключение", 1)
    add_para(doc, "Теоретически компилятор является последовательностью преобразований: текст программы постепенно превращается в токены, дерево, проверенное дерево, промежуточное представление и объектный код. Такая архитектура удобна для отладки, потому что каждая фаза имеет понятные входные и выходные данные.")
    add_para(doc, "В результате работы разработан мини-компилятор учебного языка программирования для варианта 11. Проект реализует классическую архитектуру компилятора: лексический анализ, синтаксический анализ, построение AST, семантическую проверку и генерацию кода.")
    add_para(doc, "Особое требование варианта — тернарный оператор cond ? a : b — реализовано полностью: семантически проверяются тип условия и совместимость ветвей, а в LLVM IR создаются basic blocks и PHI node. Генерация объектного файла выполнена через API LLVM 14 и legacy::PassManager, что соответствует ограничениям задания.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "13. Приложения", 1)
    appendices = [
        ("Приложение A. Полный листинг src/lexer.l", "src/lexer.l"),
        ("Приложение B. Полный листинг src/parser.y", "src/parser.y"),
        ("Приложение C. Полный листинг src/sema.hpp", "src/sema.hpp"),
        ("Приложение D. Полный листинг src/sema.cpp", "src/sema.cpp"),
        ("Приложение E. Полный листинг src/codegen.hpp", "src/codegen.hpp"),
        ("Приложение F. Полный листинг src/codegen.cpp", "src/codegen.cpp"),
        ("Приложение G. Полный листинг src/driver.cpp", "src/driver.cpp"),
        ("Приложение H. CMakeLists.txt", "CMakeLists.txt"),
        ("Приложение I. runtime/main.c", "runtime/main.c"),
    ]
    for title, path in appendices:
        add_heading(doc, title, 2)
        add_code(doc, read(path))

    try:
        doc.save(OUT)
    except PermissionError:
        fallback = OUT.with_name(OUT.stem + "_обновлённый.docx")
        doc.save(fallback)


def make_guide() -> None:
    if GUIDE.exists():
        return
    GUIDE.write_text(
        """# Как сделать скриншоты для отчёта

Все команды выполнять в WSL Ubuntu.

## 1. Структура проекта

```bash
cd "/mnt/c/Users/minae/Documents/New project"
ls
find src runtime tests -maxdepth 1 -type f | sort
```

Вставить в раздел 12, место «Структура проекта».

## 2. Успешная сборка

```bash
cd "/mnt/c/Users/minae/Documents/New project"
mkdir -p build
cd build
cmake ..
cmake --build .
```

На скриншоте должны быть видны `Found LLVM 14.0.0` и `[100%] Built target mini_cc`.

## 3. Запуск основного примера

```bash
./mini_cc ../tests/test.mc -o generated.o --emit-ir generated.ll
clang ../runtime/main.c generated.o -o app
./app
```

На скриншоте должен быть виден вывод `14`.

## 4. Проверка PHI для ternary

```bash
grep -n phi generated.ll
```

На скриншоте должны быть видны строки с `phi`.

## 5. Семантические ошибки

```bash
./mini_cc ../tests/08_type_error.mc -o bad.o --emit-ir bad.ll
./mini_cc ../tests/09_undeclared_variable.mc -o bad.o --emit-ir bad.ll
```

На скриншоте должны быть видны `semantic error` и координаты строки/столбца.

## 6. Short-circuit

```bash
./mini_cc ../tests/10_short_circuit.mc -o sc.o --emit-ir sc.ll
clang ../runtime/main.c sc.o -o sc_app
./sc_app
```

Вывод должен быть `123`. Это показывает, что правая часть `&&` не вычисляется, если левая часть ложная.

## 7. Полный прогон положительных тестов

```bash
cd "/mnt/c/Users/minae/Documents/New project/build"
for t in ../tests/01_return_literal.mc ../tests/02_arithmetic.mc ../tests/03_if_else.mc ../tests/04_for.mc ../tests/05_ternary_true.mc ../tests/06_ternary_false.mc ../tests/07_nested_ternary.mc ../tests/10_short_circuit.mc ../tests/11_function_call.mc ../tests/test.mc; do
  base=$(basename "$t" .mc)
  ./mini_cc "$t" -o "$base.o" --emit-ir "$base.ll"
  clang ../runtime/main.c "$base.o" -o "$base.app"
  printf "%s: " "$base"
  ./"$base.app"
done
```

Ожидаемый вывод:

```text
01_return_literal: 42
02_arithmetic: 35
03_if_else: 11
04_for: 45
05_ternary_true: 100
06_ternary_false: 200
07_nested_ternary: 1
10_short_circuit: 123
11_function_call: 15
test: 14
```
""",
        encoding="utf-8",
    )


if __name__ == "__main__":
    make_report()
    make_guide()
    print(OUT)
    print(GUIDE)
