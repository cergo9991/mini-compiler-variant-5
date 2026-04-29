from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\minae\Downloads\Telegram Desktop\Отчёт_исправленный.docx")
OUT = ROOT / "Отчёт_мини_компилятор_вариант_5_по_образцу.docx"
DIAGRAM = ROOT / "diagram_grammar_variant_5_railroad.png"
IMG_DIR = ROOT / "extracted_variant5_images"


def read(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def fnt(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
        "arialbd.ttf" if bold else "arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_center(draw, box, text, font, fill=(0, 0, 0), spacing=2):
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    x = box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2
    y = box[1] + (box[3] - box[1] - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=spacing, align="center")


def arrow(draw, start, end):
    draw.line([start, end], fill=(0, 0, 0), width=2)
    x, y = end
    draw.polygon([(x, y), (x - 10, y - 5), (x - 10, y + 5)], fill=(0, 0, 0))


def box(draw, xy, text, kind="nonterminal", font=None):
    font = font or fnt(15)
    if kind == "terminal":
        fill, outline, color = (255, 255, 255), (170, 30, 30), (175, 0, 0)
    elif kind == "variant":
        fill, outline, color = (232, 248, 238), (40, 125, 65), (20, 95, 45)
    else:
        fill, outline, color = (224, 238, 255), (65, 95, 178), (25, 58, 145)
    draw.rounded_rectangle(xy, radius=7, fill=fill, outline=outline, width=2)
    text_center(draw, xy, text, font, color)


def rule_panel(draw, xy, title, parts, note=""):
    draw.rounded_rectangle(xy, radius=4, fill=(255, 255, 255), outline=(0, 0, 0), width=2)
    draw.text((xy[0] + 10, xy[1] + 8), f"{title} =", font=fnt(14, True), fill=(0, 0, 0))
    y = xy[1] + 52
    arrow(draw, (xy[0] + 18, y), (xy[0] + 60, y))
    x = xy[0] + 70
    for label, kind in parts:
        w = max(54, int(draw.textlength(label, font=fnt(13))) + 22)
        if x + w + 20 > xy[2] - 15:
            x = xy[0] + 70
            y += 42
            arrow(draw, (xy[0] + 18, y), (xy[0] + 60, y))
        box(draw, (x, y - 16, x + w, y + 16), label, kind, fnt(13))
        x += w + 12
        if label != parts[-1][0]:
            arrow(draw, (x - 8, y), (x + 20, y))
            x += 28
    if note:
        text_center(draw, (xy[0] + 8, xy[3] - 28, xy[2] - 8, xy[3] - 6), note, fnt(10), (30, 30, 30))


def make_diagram() -> None:
    img = Image.new("RGB", (2200, 1500), "white")
    draw = ImageDraw.Draw(img)
    draw.text((755, 10), "MINI_CC parser grammar", font=fnt(54, True), fill=(0, 0, 0))
    draw.text((880, 72), "Railroad diagram generated from parser.y", font=fnt(22, True), fill=(20, 20, 20))

    x1, x2, x3, x4 = 25, 630, 1235, 1810
    y = 130
    rule_panel(draw, (x1, y, x1 + 585, y + 95), "program", [("function_list", "nonterminal")])
    rule_panel(draw, (x1, y + 105, x1 + 585, y + 215), "function_list", [("function", "nonterminal")], "One or more function declarations")
    rule_panel(draw, (x1, y + 230, x1 + 585, y + 350), "function", [
        ("int", "terminal"), ("IDENT", "terminal"), ("(", "terminal"), ("param_list_opt", "nonterminal"),
        (")", "terminal"), ("block", "nonterminal")
    ], "C-style function declaration")
    rule_panel(draw, (x1, y + 365, x1 + 585, y + 475), "param_list_opt", [("empty", "nonterminal"), ("param_list", "nonterminal")])
    rule_panel(draw, (x1, y + 490, x1 + 585, y + 620), "param_list", [
        ("int", "terminal"), ("IDENT", "terminal"), (",", "terminal"), ("int", "terminal"), ("IDENT", "terminal")
    ], "Zero or more parameters separated by commas")
    rule_panel(draw, (x1, y + 635, x1 + 585, y + 765), "var_decl", [
        ("int", "terminal"), ("IDENT", "terminal"), ("=", "terminal"), ("expression", "nonterminal")
    ], "Variable declaration")

    rule_panel(draw, (x2, y, x2 + 585, y + 95), "block", [("{", "terminal"), ("stmt_list", "nonterminal"), ("}", "terminal")])
    rule_panel(draw, (x2, y + 105, x2 + 585, y + 215), "stmt_list", [("empty", "nonterminal"), ("statement", "nonterminal")], "Zero or more statements")
    rule_panel(draw, (x2, y + 230, x2 + 585, y + 475), "statement", [
        ("var_decl", "nonterminal"), (";", "terminal"), ("assign", "nonterminal"), (";", "terminal"),
        ("expression", "nonterminal"), (";", "terminal"), ("return", "terminal"), ("expression", "nonterminal"),
        (";", "terminal"), ("block", "nonterminal"), ("if", "terminal"), ("for", "terminal"), ("do_while", "variant")
    ])
    rule_panel(draw, (x2, y + 490, x2 + 585, y + 620), "for_init", [("empty", "nonterminal"), ("var_decl", "nonterminal"), ("expression", "nonterminal")])
    rule_panel(draw, (x2, y + 635, x2 + 585, y + 765), "assign", [("IDENT", "terminal"), ("=", "terminal"), ("assignment", "nonterminal")], "assignment is an expression")

    rule_panel(draw, (x3, y, x3 + 555, y + 105), "argument_list_opt", [("empty", "nonterminal"), ("argument_list", "nonterminal")])
    rule_panel(draw, (x3, y + 120, x3 + 555, y + 235), "argument_list", [
        ("expression", "nonterminal"), (",", "terminal"), ("expression", "nonterminal")
    ], "One or more expressions separated by commas")
    rule_panel(draw, (x3, y + 250, x3 + 555, y + 430), "do_while", [
        ("do", "terminal"), ("block", "nonterminal"), ("while", "terminal"), ("(", "terminal"),
        ("expression", "nonterminal"), (")", "terminal"), (";", "terminal")
    ], "Variant 5: post-condition loop")

    expr = (x3, y + 445, x3 + 555, y + 765)
    draw.rounded_rectangle(expr, radius=4, fill=(255, 255, 255), outline=(0, 0, 0), width=2)
    draw.text((expr[0] + 10, expr[1] + 8), "expression =", font=fnt(14, True), fill=(0, 0, 0))
    expr_lines = [
        ("A. Primary expressions", ["INTEGER", "IDENT", "call", "( expression )"]),
        ("B. Function call", ["IDENT", "(", "argument_list_opt", ")"]),
        ("C. Unary expressions", ["!", "-", "expression"]),
        ("D. Binary expressions", ["*", "/", "%", "+", "-", "<", "<=", ">", ">=", "==", "!=", "&&", "||"]),
    ]
    yy = expr[1] + 40
    for title, vals in expr_lines:
        draw.text((expr[0] + 18, yy), title, font=fnt(11, True), fill=(0, 0, 0))
        yy += 25
        xx = expr[0] + 35
        for val in vals:
            kind = "terminal" if val in {"INTEGER", "IDENT", "!", "-", "*", "/", "%", "+", "<", "<=", ">", ">=", "==", "!=", "&&", "||", "(", ")"} else "nonterminal"
            w = max(52, int(draw.textlength(val, font=fnt(10))) + 20)
            if xx + w > expr[2] - 25:
                xx = expr[0] + 35
                yy += 34
            box(draw, (xx, yy, xx + w, yy + 26), val, kind, fnt(10))
            xx += w + 9
        yy += 38

    prec = (x4, y + 180, x4 + 360, y + 690)
    draw.rounded_rectangle(prec, radius=4, fill=(255, 255, 255), outline=(0, 0, 0), width=2)
    draw.text((prec[0] + 55, prec[1] + 25), "Operator precedence", font=fnt(22, True), fill=(0, 0, 0))
    draw.text((prec[0] + 115, prec[1] + 55), "(lowest to highest)", font=fnt(13), fill=(0, 0, 0))
    rows = [
        ("1", "=", "right-associative"),
        ("2", "||", "left-associative"),
        ("3", "&&", "left-associative"),
        ("4", "==, !=", "left-associative"),
        ("5", "<, <=, >, >=", "left-associative"),
        ("6", "+, -", "left-associative"),
        ("7", "*, /, %", "left-associative"),
        ("8", "unary ! and -", "right-associative"),
    ]
    ty = prec[1] + 95
    for n, op, assoc in rows:
        draw.rectangle((prec[0] + 15, ty - 8, prec[2] - 15, ty + 42), outline=(120, 120, 120), width=1)
        draw.text((prec[0] + 30, ty + 5), n, font=fnt(14), fill=(0, 0, 0))
        draw.text((prec[0] + 95, ty + 5), op, font=fnt(14, True), fill=(175, 0, 0))
        draw.text((prec[0] + 205, ty + 2), assoc, font=fnt(11), fill=(0, 0, 0))
        ty += 50

    token_panel = (25, 1180, 2170, 1460)
    draw.rounded_rectangle(token_panel, radius=6, fill=(255, 255, 255), outline=(0, 0, 0), width=2)
    draw.text((1000, 1190), "Parser tokens", font=fnt(22, True), fill=(0, 0, 0))
    groups = [
        ("Keywords", ["int", "if", "else", "for", "do", "while", "return"]),
        ("Identifiers and literals", ["IDENT", "INTEGER"]),
        ("Operators", ["=", "+", "-", "*", "/", "%", "==", "!=", "<", "<=", ">", ">=", "&&", "||", "!"]),
        ("Separators", [",", ";", "(", ")", "{", "}"]),
        ("Legend", ["blue box = nonterminal", "white box with red text = terminal", "green box = variant 5"]),
    ]
    gx = [45, 610, 945, 1515, 1760]
    widths = [535, 300, 540, 215, 390]
    for (name, vals), x, w in zip(groups, gx, widths):
        draw.rounded_rectangle((x, 1245, x + w, 1435), radius=5, outline=(0, 0, 0), width=2)
        draw.text((x + 25, 1260), name, font=fnt(13, True), fill=(0, 0, 0))
        xx, yy = x + 25, 1300
        for val in vals:
            kind = "terminal"
            if "nonterminal" in val:
                kind = "nonterminal"
            if "variant" in val:
                kind = "variant"
            bw = max(54, int(draw.textlength(val, font=fnt(11))) + 20)
            if xx + bw > x + w - 20:
                xx = x + 25
                yy += 38
            box(draw, (xx, yy, xx + bw, yy + 28), val, kind, fnt(11))
            xx += bw + 10

    img.save(DIAGRAM)


def clear_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = None
    if body[-1].tag == qn("w:sectPr"):
        sect_pr = deepcopy(body[-1])
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def style_exists(doc: Document, name: str) -> bool:
    return any(style.name == name for style in doc.styles)


def p(doc: Document, text="", style=None, align=None):
    para = doc.add_paragraph(style=style if style_exists(doc, style or "") else None)
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(6)
    if style is None or style == "Normal":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align is None else align
        para.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return para


def heading(doc: Document, text: str, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.name = "Times New Roman"
    return para


def caption(doc: Document, text: str):
    para = p(doc, text, "Caption Small" if style_exists(doc, "Caption Small") else None, WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.first_line_indent = Cm(0)
    for run in para.runs:
        run.font.size = Pt(10)
        run.italic = True
    return para


def code(doc: Document, text: str, max_lines: int | None = None, numbered=False):
    lines = text.rstrip().splitlines()
    if numbered:
        lines = [f"{i + 1:4d} | {line}" for i, line in enumerate(lines)]
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines] + ["     | ... фрагмент сокращен, полный файл находится в проекте"]
    para = doc.add_paragraph(style="Code Block" if style_exists(doc, "Code Block") else None)
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0.35)
    run = para.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)
    return para


def table_box(doc: Document, title: str, body: str | None = None, image: Path | None = None, width_cm=15.1):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F7F7")
    pr.append(shd)
    tp = cell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(11)
    if image and image.exists():
        ip = cell.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(str(image), width=Cm(width_cm))
    if body:
        bp = cell.add_paragraph()
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        br = bp.add_run(body)
        br.font.name = "Times New Roman"
        br.font.size = Pt(10)
    doc.add_paragraph()


def simple_table(doc: Document, rows: list[tuple[str, str]]):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.cell(0, 0).text = "Параметр"
    tbl.cell(0, 1).text = "Значение"
    for name, value in rows:
        cells = tbl.add_row().cells
        cells[0].text = name
        cells[1].text = value
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    doc.add_paragraph()


def title_page(doc: Document):
    lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "Псковский государственный университет\nПередовая инженерная школа\nгибридных технологий в станкостроении\nСоюзного государства",
        "Информатика и вычислительная техника",
        "Лингвистическое обеспечение робототехнических систем",
    ]
    for line in lines:
        para = p(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)
        para.paragraph_format.first_line_indent = Cm(0)
    p(doc, "Курсовая работа на тему:", align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].bold = True
    title = p(doc, "Разработка компилятора для учебного языка программирования с генерацией MIPS кода", align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.first_line_indent = Cm(0)
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    var = p(doc, "Вариант 5", align=WD_ALIGN_PARAGRAPH.CENTER)
    var.paragraph_format.first_line_indent = Cm(0)
    var.runs[0].bold = True
    p(doc, "\n")
    p(doc, "Выполнил\nстудент 3-го курса\nгр. 0483-04\nМинаев А.Н.", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p(doc, "Проверила\nТимошевская О. Ю.", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p(doc, "\n\n")
    p(doc, "Псков 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def toc(doc: Document):
    p(doc, "СОДЕРЖАНИЕ", align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].bold = True
    items = [
        ("toc 1", "1. Задание\t4"),
        ("toc 1", "2. Лексическая спецификация\t6"),
        ("toc 1", "3. Диаграмма разработанных правил грамматики\t8"),
        ("toc 1", "4. Листинг кода лексического и синтаксического анализаторов\t10"),
        ("toc 1", "5. Описание семантических правил языка\t12"),
        ("toc 1", "6. Листинг кода семантического анализатора\t14"),
        ("toc 1", "7. Описание процесса перевода AST в LLVM IR и далее в объектный файл\t15"),
        ("toc 1", "8. Листинг кода генератора кода\t18"),
        ("toc 1", "9. Тестовые наборы программ с пояснениями\t20"),
        ("toc 2", "Дополнительная проверка do/while при разных значениях\t23"),
        ("toc 1", "10. Описание процесса использования компилятора\t27"),
        ("toc 1", "11. Результаты тестирования\t30"),
        ("toc 1", "12. Заключение\t32"),
        ("toc 1", "13. Приложения\t35"),
    ]
    for style, text in items:
        p(doc, text, style if style_exists(doc, style) else None, WD_ALIGN_PARAGRAPH.LEFT).paragraph_format.first_line_indent = Cm(0)
    doc.add_page_break()


def build_report() -> None:
    make_diagram()
    doc = Document(TEMPLATE)
    clear_body(doc)
    title_page(doc)
    toc(doc)

    heading(doc, "1. Задание")
    p(doc, "Компилятор в рамках данной работы понимается как программа-переводчик, которая принимает на вход текст на исходном языке и выдаёт эквивалентное представление на более низком уровне. Практически любой компилятор строится как набор последовательных фаз: лексический анализ, синтаксический анализ, семантический анализ и генерация кода.")
    p(doc, "На вход разработанной системы подаётся файл с программой на учебном языке `.mc`, а на выходе формируются промежуточное представление LLVM IR и объектный файл. В данной работе все фазы реализованы отдельно: лексер написан на Flex, parser — на GNU Bison, семантическая проверка и генерация кода реализованы на C++17 с использованием LLVM 14.")
    p(doc, "В рамках курсовой работы реализован учебный мини-компилятор для варианта 5. Вариант требует C-style syntax, конструкцию `do/while` и генерацию объектного файла для архитектуры MIPS.")
    for item in [
        "• минимальный набор конструкций: арифметические и логические операции, условные операторы `if/else`, цикл `for`, функции и тип `int`;",
        "• стиль синтаксиса: C-style, функции имеют вид `int name(int arg) { ... }`;",
        "• дополнительная конструкция варианта: цикл с постусловием `do { ... } while (condition);`;",
        "• целевая архитектура: MIPS с целевой строкой `mipsel-unknown-linux-gnu`;",
        "• практический результат: генерация файлов `.ll` и `.o`, а при наличии MIPS toolchain — компоновка с `runtime/main.c` и запуск через `qemu-mipsel`.",
    ]:
        p(doc, item)
    simple_table(doc, [
        ("Вариант", "5"),
        ("Конструкция", "Do/While"),
        ("Архитектура", "MIPS"),
        ("Синтаксис", "C-style syntax"),
        ("Цель LLVM", "mipsel-unknown-linux-gnu"),
    ])

    heading(doc, "2. Лексическая спецификация")
    p(doc, "Лексический анализатор предназначен для преобразования непрерывного потока символов в последовательность токенов. На вход он получает исходный текст программы, а на выход передаёт синтаксическому анализатору коды токенов и связанные с ними значения.")
    p(doc, "В данном проекте лексическая фаза реализована в Flex. Лексер распознаёт ключевые слова, идентификаторы, целочисленные литералы, комментарии, операторы и разделители. Для варианта 5 добавлены ключевые слова `do` и `while`.")
    for item in [
        "• ключевые слова: `int`, `return`, `if`, `else`, `for`, `do`, `while`;",
        "• идентификаторы: шаблон [A-Za-z_][A-Za-z0-9_]*;",
        "• целочисленные литералы: последовательность цифр `[0-9]+`;",
        "• арифметические, логические и сравнительные операторы;",
        "• разделители: запятая, точка с запятой, круглые и фигурные скобки;",
        "• однострочные и многострочные комментарии, не влияющие на синтаксический анализ.",
    ]:
        p(doc, item)

    heading(doc, "3. Диаграмма разработанных правил грамматики")
    p(doc, "Синтаксический анализ определяет, как из потока токенов строится корректная структура программы. Его входом является последовательность токенов от лексического анализатора, а выходом — абстрактное синтаксическое дерево, используемое последующими фазами.")
    p(doc, "Грамматика реализована в GNU Bison и разбита на нетерминалы верхнего уровня: программа, список функций, объявление функции, блок, список операторов, оператор и выражение. Отдельное правило варианта 5 задаёт цикл do/while.")
    table_box(doc, "Диаграмма грамматики", image=DIAGRAM, width_cm=16.5)

    heading(doc, "4. Листинг кода лексического и синтаксического анализаторов")
    p(doc, "В основной части отчёта приведены ключевые фрагменты кода, показывающие принцип работы лексера и parser. Полные листинги вынесены в приложения, чтобы основной текст оставался сосредоточен на устройстве компилятора.")
    caption(doc, "Листинг 4.1 — фрагмент файла src/lexer.l")
    code(doc, read("src/lexer.l"), max_lines=90)
    caption(doc, "Листинг 4.2 — фрагмент файла src/parser.y")
    code(doc, read("src/parser.y"), max_lines=120)

    heading(doc, "5. Описание семантических правил языка")
    p(doc, "Семантический анализ выполняет содержательную проверку программы после того, как её структура уже распознана. На вход этой фазе передаётся AST, а на выходе формируется либо подтверждение корректности программы, либо список диагностических сообщений.")
    for item in [
        "1. в одной области видимости запрещено повторное объявление переменной;",
        "2. использование необъявленного имени считается ошибкой;",
        "3. поддерживается тип `int`; тип `bool` используется как результат сравнений и логических выражений;",
        "4. арифметические операции разрешены только для `int`; сравнения возвращают `bool`;",
        "5. условия `if`, `for` и `do/while` обязаны иметь тип `bool`;",
        "6. вызовы функций проверяются по числу аргументов и их типам;",
        "7. для функции `compiled_fn` проверяется сигнатура `int compiled_fn(int arg)`;",
        "8. цикл `do/while` анализирует тело как отдельный блок, а затем проверяет условие после тела.",
    ]:
        p(doc, item)
    p(doc, "Особенность do/while состоит в том, что тело выполняется минимум один раз, а условие проверяется после выполнения тела. В семантическом анализаторе это отражено отдельной веткой для узла `DoWhileStmt`.")

    heading(doc, "6. Листинг кода семантического анализатора")
    p(doc, "Ниже оставлен ключевой фрагмент семантического анализатора: регистрация функций, работа с областями видимости, проверка типов выражений и отдельная проверка условия do/while.")
    caption(doc, "Листинг 6.1 — ключевой фрагмент src/sema.cpp")
    code(doc, read("src/sema.cpp"), max_lines=165)
    p(doc, "Продолжение полного листинга см. в приложении C.")

    heading(doc, "7. Описание процесса перевода абстрактного синтаксического дерева (AST) в промежуточное представление LLVM (LLVM IR) и далее в объектный файл")
    p(doc, "Генерация кода переводит проверенное AST в форму, пригодную для дальнейшей машинной обработки. В работе сначала строится LLVM IR, затем модуль передаётся объекту TargetMachine, который генерирует объектный файл для MIPS.")
    for item in [
        "1. создаются объекты `LLVMContext`, `Module` и `IRBuilder`;",
        "2. для каждой функции объявляется сигнатура и создаётся входной блок `entry`;",
        "3. локальные переменные размещаются через `alloca`, параметры функции сохраняются в локальные слоты;",
        "4. конструкции `if/else` и `for` транслируются в набор basic blocks с явными переходами;",
        "5. логические `&&` и `||` используют short-circuit и PHI-узлы;",
        "6. конструкция `do/while` создаёт блоки `do.body`, `do.cond` и `do.end`;",
        "7. объектный файл формируется для target triple `mipsel-unknown-linux-gnu`.",
    ]:
        p(doc, item)
    table_box(doc, "Сгенерированный LLVM IR и MIPS object", "На скриншоте показано, что объектный файл имеет формат MIPS ELF, а LLVM IR содержит блоки do.body, do.cond и do.end.", IMG_DIR / "image3.png", 15.1)
    table_box(doc, "Проверка блоков do/while в IR", "Переход `br i1` возвращает управление в тело при истинном условии и выходит в do.end при ложном.", IMG_DIR / "image4.png", 15.1)

    heading(doc, "8. Листинг кода генератора кода")
    p(doc, "Генератор кода реализует обход AST и создание инструкций LLVM. Литералы превращаются в константы, переменные работают через `alloca`, `load` и `store`, а управляющие конструкции создают basic blocks.")
    p(doc, "Для do/while сначала создаётся безусловный переход в `do.body`. После генерации тела создаётся переход в `do.cond`, где вычисляется условие. Если условие истинно, управление возвращается в тело; если ложно — переходит в `do.end`.")
    caption(doc, "Листинг 8.1 — фрагмент src/codegen.cpp")
    code(doc, read("src/codegen.cpp"), max_lines=180)

    heading(doc, "9. Тестовые наборы программ с пояснениями")
    p(doc, "Тестирование проверяет весь конвейер: исходный файл `.mc`, разбор, семантическую проверку, генерацию LLVM IR и создание MIPS object file.")
    tests = [
        ("01_return_literal.mc", "возврат целочисленного литерала"),
        ("02_arithmetic.mc", "арифметические операции и приоритеты"),
        ("03_if_else.mc", "условная конструкция if/else"),
        ("04_for.mc", "цикл for из базового набора"),
        ("05_do_while_sum.mc", "суммирование в do/while"),
        ("06_do_while_at_least_once.mc", "проверка выполнения тела минимум один раз"),
        ("07_do_while_power.mc", "несколько итераций do/while"),
        ("08_type_error.mc", "ожидаемая ошибка типов"),
        ("09_undeclared_variable.mc", "ожидаемая ошибка необъявленной переменной"),
        ("10_short_circuit.mc", "short-circuit для &&"),
        ("11_function_call.mc", "вызов пользовательской функции"),
    ]
    for name, desc in tests:
        p(doc, f"• `{name}` — {desc}.")
    caption(doc, "Листинг 9.1 — тест do/while")
    code(doc, read("tests/05_do_while_sum.mc"))
    heading(doc, "Дополнительная проверка do/while при разных значениях", 2)
    p(doc, "Для ручной проверки удобно менять аргумент в `runtime/main.c` и повторно запускать связывание через MIPS toolchain. При аргументе 10 тест `05_do_while_sum.mc` возвращает сумму чисел от 0 до 9, то есть 45.")
    heading(doc, "Команда 9.2 — запуск do_while_sum", 3)
    code(doc, "build_variant5/mini_cc tests/05_do_while_sum.mc -o build_variant5/05_do_while_sum.o --emit-ir build_variant5/05_do_while_sum.ll")
    heading(doc, "Команда 9.3 — проверка MIPS object", 3)
    code(doc, "file build_variant5/05_do_while_sum.o\ngrep -n -E 'do\\.body|do\\.cond|do\\.end|br i1' build_variant5/05_do_while_sum.ll")

    heading(doc, "10. Описание процесса использования компилятора")
    p(doc, "Практическое использование компилятора сводится к трём основным действиям: сборке проекта, компиляции исходной программы и проверке полученного объектного файла. Пользовательский сценарий построен вокруг CMake, исполняемого файла `mini_cc`, LLVM IR и MIPS object output.")
    caption(doc, "Команда 10.1")
    code(doc, "cmake -S . -B build_variant5\ncmake --build build_variant5")
    table_box(doc, "Сборка проекта", "На скриншоте представлен вывод CMake и успешная сборка цели mini_cc.", IMG_DIR / "image2.png", 15.1)
    caption(doc, "Команда 10.2")
    code(doc, "build_variant5/mini_cc tests/05_do_while_sum.mc -o build_variant5/05_do_while_sum.o --emit-ir build_variant5/05_do_while_sum.ll")
    caption(doc, "Команда 10.3")
    code(doc, "file build_variant5/05_do_while_sum.o")

    heading(doc, "11. Результаты тестирования")
    p(doc, "Результаты тестирования показывают, что реализованный компилятор соответствует требованиям варианта 5. Проект успешно собирается, создаёт LLVM IR, генерирует MIPS object file и обнаруживает семантические ошибки до этапа генерации кода.")
    for item in [
        "• `05_do_while_sum`: сгенерирован MIPS object file;",
        "• LLVM IR содержит блоки `do.body`, `do.cond` и `do.end`;",
        "• `08_type_error`: ошибка типа инициализатора корректно диагностируется;",
        "• `09_undeclared_variable`: использование необъявленной переменной корректно диагностируется;",
        "• положительные тесты компилируются в `.ll` и `.o` без ошибок.",
    ]:
        p(doc, item)
    table_box(doc, "Пример семантической ошибки", "На скриншоте показана ошибка типов: переменная `x` имеет тип int, но инициализатор является bool-выражением.", IMG_DIR / "image5.png", 15.1)
    caption(doc, "Листинг 11.1 — ожидаемый вывод проверки")
    code(doc, dedent("""\
        build_variant5/report_check.o: ELF 32-bit LSB relocatable, MIPS, MIPS32 version 1 (SYSV), not stripped
        target triple = "mipsel-unknown-linux-gnu"
        do.body:
        do.cond:
          br i1 %lttmp, label %do.body, label %do.end
        semantic error: line 2, column 5: initializer for variable 'x' has type bool, expected int
        """))

    heading(doc, "12. Заключение")
    p(doc, "В ходе выполнения курсовой работы был разработан учебный мини-компилятор для языка программирования с C-style синтаксисом и генерацией объектного файла под архитектуру MIPS. Реализованный проект охватывает основные этапы компиляции: лексический анализ, синтаксический анализ, построение AST, семантическую проверку и генерацию кода.")
    p(doc, "Во входной части компилятора реализован лексический анализатор на Flex. Он распознаёт ключевые слова языка, идентификаторы, целочисленные литералы, операторы и разделители. Синтаксический анализ реализован с использованием GNU Bison; в грамматике описаны функции, блоки, объявления переменных, выражения, условные конструкции, цикл for и конструкция do/while варианта 5.")
    p(doc, "Семантический анализ проверяет области видимости, повторные объявления, использование необъявленных имён, типы выражений, корректность условий и сигнатуру обязательной функции `compiled_fn`. Модуль генерации кода построен на базе LLVM и формирует MIPS object file для triple `mipsel-unknown-linux-gnu`.")
    p(doc, "Таким образом, задача варианта 5 выполнена: конструкция do/while добавлена во все фазы компилятора, а результатом работы является объектный файл для MIPS.")

    heading(doc, "13. Приложения")
    p(doc, "Приложения содержат полные листинги исходных файлов и служат справочным дополнением к основной части отчёта.")
    for title, path in [
        ("Приложение A. Полный листинг src/lexer.l", "src/lexer.l"),
        ("Приложение B. Полный листинг src/parser.y", "src/parser.y"),
        ("Приложение C. Полный листинг src/sema.cpp", "src/sema.cpp"),
        ("Приложение D. Полный листинг src/codegen.cpp", "src/codegen.cpp"),
    ]:
        heading(doc, title, 2)
        caption(doc, title.split(". ")[0])
        code(doc, read(path), numbered=True)

    doc.save(OUT)
    print(OUT)
    print(DIAGRAM)


if __name__ == "__main__":
    build_report()
